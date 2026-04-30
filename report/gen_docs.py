from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── TECHNICAL REPORT ──────────────────────────────────────────────────────────

doc = Document()

t = doc.add_heading('Technical Report', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Expense Management System with Ranking Logic')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

info = doc.add_paragraph('Course: Web Programming  |  Domain: Expense Management  |  Logic: Ranking System  |  Complexity: Intermediate')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 1
doc.add_heading('1. Problem Definition', 1)
doc.add_paragraph(
    'Managing personal finances is a recurring challenge for individuals. Most people spend without visibility '
    'into which categories consume the most money. This system solves that by:'
)
doc.add_paragraph('Allowing users to log expenses with categories and amounts.', style='List Bullet')
doc.add_paragraph('Automatically ranking spending categories from highest to lowest total spend.', style='List Bullet')
doc.add_paragraph('Surfacing the top individual transactions so users know where large single costs occur.', style='List Bullet')

# 2
doc.add_heading('2. System Architecture', 1)
doc.add_paragraph('Request Flow:')
doc.add_paragraph('User interacts with React UI (login, add expense, view rankings).', style='List Number')
doc.add_paragraph('React sends HTTP request to Express API with JWT in Authorization header.', style='List Number')
doc.add_paragraph('Express middleware verifies JWT, validates input, then queries MongoDB via Mongoose.', style='List Number')
doc.add_paragraph('MongoDB returns data; Express sends JSON response back to React.', style='List Number')
doc.add_paragraph('React updates state and re-renders the UI.', style='List Number')

doc.add_paragraph()
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
hdr[2].text = 'Responsibility'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
rows_data = [
    ('M', 'MongoDB Atlas', 'Persist users and expenses'),
    ('E', 'Express.js', 'REST API, routing, middleware'),
    ('R', 'React + Vite', 'UI rendering, state, routing'),
    ('N', 'Node.js', 'Runtime for Express'),
]
for i, (a, b, c) in enumerate(rows_data, 1):
    r = table.rows[i].cells
    r[0].text = a
    r[1].text = b
    r[2].text = c

# 3
doc.add_heading('3. Database Design', 1)
doc.add_heading('Collections', 2)

p = doc.add_paragraph()
p.add_run('Collection 1: users').bold = True
doc.add_paragraph(
    '_id: ObjectId\n'
    'name: String (required, max 100)\n'
    'email: String (required, unique, lowercase)\n'
    'password: String (bcrypt hashed)\n'
    'createdAt / updatedAt: Date'
).style = 'No Spacing'

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Collection 2: expenses').bold = True
doc.add_paragraph(
    '_id: ObjectId\n'
    'userId: ObjectId  =>  ref: User  (FOREIGN KEY REFERENCE)\n'
    'title: String (required, max 200)\n'
    'amount: Number (required, min 0.01)\n'
    'category: String (enum: Food | Transport | Housing | Entertainment | Healthcare | Education | Shopping | Utilities | Other)\n'
    'date: Date (required)\n'
    'description: String (optional)'
).style = 'No Spacing'

doc.add_heading('Relationship: Referencing vs Embedding', 2)
doc.add_paragraph(
    'Choice: Referencing. The expenses collection stores a userId field (ObjectId) that points to the users collection.'
)
p_just = doc.add_paragraph()
p_just.add_run('Justification:').bold = True
doc.add_paragraph(
    'A user can accumulate thousands of expenses. Embedding inside the user document would cause unbounded growth, '
    'eventually exceeding MongoDB\'s 16MB document size limit.',
    style='List Bullet'
)
doc.add_paragraph(
    'Referencing lets expenses be queried, filtered, sorted, and aggregated independently without loading the full user object.',
    style='List Bullet'
)
doc.add_paragraph(
    'Read patterns always query expenses filtered by userId, making a separate collection more efficient.',
    style='List Bullet'
)

doc.add_heading('Indexes', 2)
doc.add_paragraph('{ userId: 1, date: -1 }  — covers date-sorted queries', style='List Bullet')
doc.add_paragraph('{ userId: 1, category: 1 }  — covers category-filtered queries', style='List Bullet')
doc.add_paragraph('{ userId: 1, amount: -1 }  — covers amount-sorted ranking queries', style='List Bullet')

# 4
doc.add_heading('4. Core Logic Explanation - Ranking System', 1)
doc.add_paragraph(
    'The ranking system is the primary intelligence of the application. It determines which spending categories '
    'and transactions are most significant for a user within a selected time period.'
)
doc.add_heading('Step-by-Step Category Ranking', 2)
doc.add_paragraph(
    'Step 1 - Filter: MongoDB $match stage selects only the authenticated user\'s expenses within the chosen date range.',
    style='List Number'
)
doc.add_paragraph(
    'Step 2 - Group & Aggregate: $group by category, computing totalAmount ($sum), count ($sum: 1), '
    'avgAmount ($avg), and maxAmount ($max) for each group.',
    style='List Number'
)
doc.add_paragraph(
    'Step 3 - Sort: $sort by totalAmount descending. The category with the highest total comes first.',
    style='List Number'
)
doc.add_paragraph(
    'Step 4 - Rank: JavaScript .map() assigns rank numbers. rank 1 = highest spending category.',
    style='List Number'
)
doc.add_paragraph(
    'Step 5 - Percentage: percentage = (categoryTotal / grandTotal) x 100, showing relative spending weight.',
    style='List Number'
)
doc.add_paragraph()
doc.add_paragraph('Example output:')
doc.add_paragraph('Rank 1: Housing -- PKR 25,000 (42%)', style='List Bullet')
doc.add_paragraph('Rank 2: Food -- PKR 18,000 (30%)', style='List Bullet')
doc.add_paragraph('Rank 3: Transport -- PKR 8,000 (13%)', style='List Bullet')

# 5
doc.add_heading('5. Query Explanation', 1)
doc.add_heading('Query 1 - Category Ranking Aggregation (MongoDB Pipeline)', 2)
code1 = doc.add_paragraph(
    'Expense.aggregate([\n'
    '  { $match: { userId: ObjectId(userId), date: { $gte: startDate } } },\n'
    '  { $group: {\n'
    '    _id: "$category",\n'
    '    totalAmount: { $sum: "$amount" },\n'
    '    count: { $sum: 1 },\n'
    '    avgAmount: { $avg: "$amount" }\n'
    '  }},\n'
    '  { $sort: { totalAmount: -1 } }\n'
    '])'
)
code1.runs[0].font.name = 'Courier New'
code1.runs[0].font.size = Pt(9)
doc.add_paragraph(
    'Purpose: Groups all user expenses by category and computes totals. The $sort stage produces the ranked order. '
    'A single pipeline replaces multiple queries and client-side sorting.'
)
doc.add_paragraph('Sample input: Food: 500, Food: 300, Transport: 900')
doc.add_paragraph('Sample output: Transport: 900 (Rank 1), Food: 800 (Rank 2)')

doc.add_heading('Query 2 - Top 10 Expenses by Amount', 2)
code2 = doc.add_paragraph(
    'Expense.aggregate([\n'
    '  { $match: { userId: ObjectId(userId) } },\n'
    '  { $sort: { amount: -1 } },\n'
    '  { $limit: 10 }\n'
    '])'
)
code2.runs[0].font.name = 'Courier New'
code2.runs[0].font.size = Pt(9)
doc.add_paragraph(
    'Purpose: Returns the 10 largest individual expenses. Uses the compound index { userId: 1, amount: -1 } '
    'for efficient lookup without a full collection scan.'
)

# 6
doc.add_heading('6. Security Analysis', 1)
sec_items = [
    ('Password Hashing (bcrypt, cost factor 12)',
     'Passwords are never stored in plaintext. bcrypt.hash(password, 12) is called in a Mongoose pre-save hook before writing to MongoDB.'),
    ('JWT Authentication',
     'All protected routes require a valid Bearer token in the Authorization header. The auth middleware calls jwt.verify(). Invalid or expired tokens return HTTP 401.'),
    ('Input Validation (express-validator)',
     'All endpoints validate request bodies. title, email, password, amount, and category are validated for type, length, and format. Invalid input returns HTTP 422 with field-level details.'),
    ('Protected Routes (Frontend + Backend)',
     'Frontend: ProtectedRoute component redirects unauthenticated users to /login. Backend: Every /api/expenses route applies the auth middleware. Users can only access their own data.'),
    ('Authorization Scoping',
     'DELETE /api/expenses/:id filters by both _id AND userId. A user cannot delete another user\'s expense even if they know the ID.'),
    ('CORS Configuration',
     'CORS only allows requests from the specific CLIENT_URL environment variable, never a wildcard *.'),
]
for title, body in sec_items:
    p = doc.add_paragraph()
    p.add_run(title + ':').bold = True
    p.add_run(' ' + body)

# 7
doc.add_heading('7. Scalability Discussion', 1)
doc.add_paragraph('Issues that would arise at 10,000 users and their fixes:')
doc.add_paragraph()
table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
h2 = table2.rows[0].cells
h2[0].text = 'Issue'
h2[1].text = 'Impact'
h2[2].text = 'Fix'
for cell in h2:
    cell.paragraphs[0].runs[0].bold = True
scale = [
    ('GET /expenses returns up to 500 docs', 'Memory spike under load', 'Add cursor-based pagination (?cursor=lastId&limit=20)'),
    ('Rankings aggregation on every request', 'CPU pressure on MongoDB', 'Cache results in Redis with 5-min TTL, invalidate on new expense'),
    ('No rate limiting on /auth/login', 'Brute force login attacks', 'Add express-rate-limit: 10 requests/min per IP'),
    ('No refresh token rotation', 'Compromised tokens valid 7 days', 'Use 15-min access tokens + refresh token rotation'),
    ('Single MongoDB connection string', 'Connection pool exhaustion', 'Configure Mongoose poolSize: 10 + Atlas auto-scaling'),
    ('No data archival strategy', 'Storage cost growth', 'Archive expenses older than 2 years to cold storage'),
]
for i, (a, b, c) in enumerate(scale, 1):
    r = table2.rows[i].cells
    r[0].text = a
    r[1].text = b
    r[2].text = c

doc.save(r'C:\Users\tariq\Downloads\Nafay_WP\report\technical-report.docx')
print('Saved technical-report.docx')

# ── HANDWRITTEN ANSWERS ───────────────────────────────────────────────────────

doc2 = Document()

h = doc2.add_heading('Handwritten Answers', 0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2 = doc2.add_paragraph('Expense Management with Ranking System')
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].bold = True
doc2.add_paragraph('Print this and write / copy the answers onto your answer sheet.').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc2.add_paragraph()

# Q1
doc2.add_heading('Q1. System Flow Diagram  (2 marks)', 1)
doc2.add_paragraph('Draw and label the following flow on paper:')
doc2.add_paragraph()
flow = doc2.add_paragraph(
    '[User Browser]\n'
    '      |\n'
    '      |  HTTP Request with JWT in Authorization header\n'
    '      v\n'
    '[React Frontend  --  Vercel]\n'
    '      |\n'
    '      |  axios.get / axios.post to backend API URL\n'
    '      v\n'
    '[Express Backend  --  Render / Railway]\n'
    '      |\n'
    '      |  1. auth middleware  =>  verify JWT\n'
    '      |  2. validate middleware  =>  check input\n'
    '      |  3. Route handler  =>  call Mongoose\n'
    '      v\n'
    '[MongoDB Atlas  --  Cloud Cluster]\n'
    '      |\n'
    '      |  Returns document(s)\n'
    '      v\n'
    '[Express]  =>  JSON Response (200 / 201 / 4xx)\n'
    '      |\n'
    '      v\n'
    '[React]  =>  setState()  =>  Re-render UI'
)
flow.runs[0].font.name = 'Courier New'
flow.runs[0].font.size = Pt(9)

doc2.add_paragraph()
p_note = doc2.add_paragraph()
p_note.add_run('Layer labels to write on diagram:').bold = True
doc2.add_paragraph('Frontend: React + React Router (deployed on Vercel)', style='List Bullet')
doc2.add_paragraph('Backend: Node.js + Express (deployed on Render)', style='List Bullet')
doc2.add_paragraph('Database: MongoDB Atlas (Cloud Cluster)', style='List Bullet')
doc2.add_paragraph('Auth: JWT Bearer Token sent in every protected request header', style='List Bullet')

# Q2
doc2.add_heading('Q2. One API Route - Step by Step  (2 marks)', 1)
p_route = doc2.add_paragraph()
p_route.add_run('Route chosen: POST /api/expenses  (Add a new expense)').bold = True
doc2.add_paragraph()

doc2.add_paragraph(
    'Step 1 - Client sends POST request to /api/expenses with JSON body:\n'
    '{ "title": "Lunch", "amount": 350, "category": "Food" }\n'
    'Authorization header: Bearer <jwt_token>',
    style='List Number'
)
doc2.add_paragraph(
    'Step 2 - auth middleware runs first. It reads req.headers.authorization, splits on "Bearer ", '
    'extracts the token, and calls jwt.verify(token, JWT_SECRET). '
    'If valid, it sets req.userId = decoded.userId and calls next(). '
    'If invalid, returns 401 INVALID_TOKEN.',
    style='List Number'
)
doc2.add_paragraph(
    'Step 3 - express-validator rules run. Checks: title is not empty, amount is a float >= 0.01, '
    'category is in the allowed enum list. '
    'If any check fails, validate middleware returns 422 VALIDATION_ERROR with field-level details.',
    style='List Number'
)
doc2.add_paragraph(
    'Step 4 - Route handler executes. Calls Expense.create({ userId: req.userId, title, amount, category, date, description }). '
    'Mongoose runs schema validation, then inserts the document into MongoDB Atlas.',
    style='List Number'
)
doc2.add_paragraph(
    'Step 5 - On success, returns 201 Created with the new expense document as JSON. '
    'On any unhandled error, Express global error handler returns 500 INTERNAL_ERROR.',
    style='List Number'
)

# Q3
doc2.add_heading('Q3. Database Relationship  (2 marks)', 1)
doc2.add_paragraph('Two collections: users and expenses.')
doc2.add_paragraph('Relationship type: One-to-Many. One user has many expenses.')
doc2.add_paragraph()
p_choice = doc2.add_paragraph()
p_choice.add_run('Choice: Referencing (not embedding)').bold = True
doc2.add_paragraph(
    'The expenses collection stores a userId field of type ObjectId that points to the _id in the users collection.'
)
doc2.add_paragraph()
p_why = doc2.add_paragraph()
p_why.add_run('Why referencing over embedding:').bold = True
doc2.add_paragraph(
    'If I embedded expenses as an array inside the user document, that document would grow without bound as the user '
    'adds more expenses. Eventually it would exceed MongoDB\'s 16MB document size limit, causing writes to fail.',
    style='List Bullet'
)
doc2.add_paragraph(
    'Every query that reads expenses (filtering by category, sorting by amount, running aggregations) would be forced '
    'to load the entire user document including all embedded expenses, which is wasteful.',
    style='List Bullet'
)
doc2.add_paragraph(
    'With referencing, expenses are a separate collection queried directly with { userId: req.userId }, '
    'keeping each document small and queries efficient.',
    style='List Bullet'
)

# Q4
doc2.add_heading('Q4. Core FinTech Logic - Ranking System  (2 marks)', 1)
doc2.add_paragraph('The ranking logic lives in GET /api/expenses/rankings in the backend (routes/expenses.js).')
doc2.add_paragraph()
p_ex = doc2.add_paragraph()
p_ex.add_run('Example trace with sample data:').bold = True
doc2.add_paragraph('User has these expenses for the month:')
doc2.add_paragraph('Food: PKR 500 + PKR 800 = total PKR 1,300', style='List Bullet')
doc2.add_paragraph('Transport: PKR 400 + PKR 600 = total PKR 1,000', style='List Bullet')
doc2.add_paragraph('Housing: PKR 15,000 = total PKR 15,000', style='List Bullet')
doc2.add_paragraph()
doc2.add_paragraph(
    'Step 1 - $match: MongoDB selects only this user\'s expenses (userId matches) for the current month.',
    style='List Number'
)
doc2.add_paragraph(
    'Step 2 - $group: Groups by category and sums amounts. Result: { Food: 1300, Transport: 1000, Housing: 15000 }',
    style='List Number'
)
doc2.add_paragraph(
    'Step 3 - $sort: Orders by totalAmount descending. Order: Housing (15000) => Food (1300) => Transport (1000)',
    style='List Number'
)
doc2.add_paragraph(
    'Step 4 - Rank assignment: JavaScript .map() adds rank numbers. Housing = Rank 1, Food = Rank 2, Transport = Rank 3.',
    style='List Number'
)
doc2.add_paragraph(
    'Step 5 - Percentage: Grand total = 17,300. Housing = 86.7%, Food = 7.5%, Transport = 5.8%. '
    'Frontend displays these as rank cards with color-coded progress bars.',
    style='List Number'
)

# Q5
doc2.add_heading('Q5. Real Security Flaw  (2 marks)', 1)
p_vuln = doc2.add_paragraph()
p_vuln.add_run('Vulnerability: Missing rate limiting on the login endpoint').bold = True
doc2.add_paragraph()
doc2.add_paragraph('Location: POST /api/auth/login (backend/routes/auth.js)')
doc2.add_paragraph()
p_flaw = doc2.add_paragraph()
p_flaw.add_run('The flaw:').bold = True
doc2.add_paragraph(
    'There is no rate limiting on the login route. An attacker can send thousands of login requests per second '
    'with different passwords, attempting to brute-force a user\'s account. Since the endpoint responds with '
    'different error messages for wrong email vs wrong password, an attacker can also enumerate which emails '
    'are registered in the system.'
)
doc2.add_paragraph()
p_att = doc2.add_paragraph()
p_att.add_run('Example attack:').bold = True
att = doc2.add_paragraph(
    'POST /api/auth/login { email: "victim@email.com", password: "pass1" } => 401\n'
    'POST /api/auth/login { email: "victim@email.com", password: "pass2" } => 401\n'
    'POST /api/auth/login { email: "victim@email.com", password: "pass3" } => 200  (found it!)\n'
    'With 1,000 requests/second, this takes seconds against a common password.'
)
att.runs[0].font.name = 'Courier New'
att.runs[0].font.size = Pt(9)
doc2.add_paragraph()
p_fix = doc2.add_paragraph()
p_fix.add_run('The fix:').bold = True
doc2.add_paragraph(' Add express-rate-limit middleware on the auth router:', style='No Spacing')
fix = doc2.add_paragraph(
    'const rateLimit = require("express-rate-limit");\n'
    'const loginLimiter = rateLimit({\n'
    '  windowMs: 15 * 60 * 1000,   // 15 minutes\n'
    '  max: 10,                      // 10 attempts per IP\n'
    '  message: { error: { code: "TOO_MANY_REQUESTS",\n'
    '    message: "Too many login attempts. Try again in 15 minutes." } }\n'
    '});\n'
    'router.post("/login", loginLimiter, [...validators], validate, handler);'
)
fix.runs[0].font.name = 'Courier New'
fix.runs[0].font.size = Pt(9)
doc2.add_paragraph(
    'This limits each IP to 10 login attempts per 15 minutes, making brute-force attacks completely impractical.'
)

doc2.save(r'C:\Users\tariq\Downloads\Nafay_WP\report\handwritten-answers.docx')
print('Saved handwritten-answers.docx')
